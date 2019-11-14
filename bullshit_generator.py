import os
import re
import random
import readJSON
from docx import Document


data = readJSON.read_json("data.json")
idiom = data["famous"]  # a 代表前面垫话，b代表后面垫话
before = data["before"]  # 在名人名言前面弄点废话
after = data['after']  # 在名人名言后面弄点废话
shit_start = data['bosh_start']
shit = data['bosh']  # 代表文章主要废话来源

title = "废话"

repeat = 2


def random_loop(l):
    global repeat
    pool = list(l) * repeat
    while True:
        random.shuffle(pool)
        for i in pool:
            yield i

            
next_shit = random_loop(shit)
next_idiom = random_loop(idiom)
next_shit_start = random_loop(shit_start)


def add_idiom():
    global next_idiom
    xx = next(next_idiom)
    xx = xx.replace("a", random.choice(before))
    xx = xx.replace("b", random.choice(after))
    return xx


def another_paragraph():
    xx = ""
    xx += "\r\n"
    xx += "    "
    return xx


if __name__ == "__main__":
    xx = input("请输入文章主题: ")
    length = int(input("字數: "))
    fn = input('文件名: ')
    eles = input('具备元素，支持多个，半角空格隔开: ').split()
    total = 0
    # tmp = str("    ")
    tmp = Document()
    tmp.add_heading(xx, 0)
    p = tmp.add_paragraph()
    while (total < length):
        conditiion = random.randint(0, 100)
        conditiion1 = random.randint(0, 100)
        if conditiion < 20:
            # tmp += another_paragraph()
            if len(p.text) > 0:
                p = tmp.add_paragraph()
        elif conditiion < 40:
            s = ''
            if conditiion1 < 40:
                s += next(next_shit_start)
                
            s += add_idiom()
            if eles:
                s = s.replace("x", random.choice(eles))
            else:
                s = s.replace("x", xx)
            total += len(s)
            # tmp += s
            p.add_run(s)
        else:
            s = ''
            if conditiion1 < 60:
                s += next(next_shit_start)
            s += next(next_shit)
            if eles:
                s = s.replace("x", random.choice(eles))
            else:
                s = s.replace("x", xx)
            total += len(s)
            # tmp += s
            p.add_run(s)
    #tmp = tmp.replace("x", xx)
    #print(tmp)
    tmp.save('{}.docx'.format(fn))
